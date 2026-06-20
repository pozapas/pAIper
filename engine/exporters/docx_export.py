"""Render a ReviewReport to a polished Word (.docx) document."""

from __future__ import annotations

import io

from ..report import ReviewReport
from .branding import GENERATED_BY, PAIPER_REPO_LABEL, PAIPER_REPO_URL, SEP, usage_suffix

# severity → (label color RGB, fill hex)
_SEV_COLOR = {
    "CRITICAL": ((0xB9, 0x1C, 0x1C), "FEE2E2"),
    "MAJOR": ((0xC2, 0x41, 0x0C), "FFEDD5"),
    "MINOR": ((0xA1, 0x62, 0x07), "FEF9C3"),
    "PRAISE": ((0x15, 0x80, 0x3D), "DCFCE7"),
}
_STATUS_MARK = {"PASS": "PASS", "FAIL": "FAIL", "UNCLEAR": "UNCLEAR", "NA": "N/A"}
_REC_COLOR = {
    "Accept": (0x15, 0x80, 0x3D),
    "Minor Revision": (0xA1, 0x62, 0x07),
    "Major Revision": (0xC2, 0x41, 0x0C),
    "Reject": (0xB9, 0x1C, 0x1C),
}


def report_to_docx(r: ReviewReport) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    def heading(text, level=1, color=(0x1E, 0x3A, 0x8A)):
        p = doc.add_heading(text, level=level)
        if p.runs:
            p.runs[0].font.color.rgb = RGBColor(*color)
        return p

    def para(text="", *, bold=False, italic=False, color=None, size=None, indent=None):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Inches(indent)
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
        if size:
            run.font.size = Pt(size)
        return p

    def shaded(text, fill="F1F5F9", color=(0x47, 0x55, 0x69), mono=True, indent=0.25):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(indent)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        pPr.append(shd)
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(*color)
        if mono:
            run.font.name = "Consolas"
        return p

    def add_hyperlink(paragraph, text, url, color="FF9F43"):
        r_id = paragraph.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        i = OxmlElement("w:i")
        rPr.append(i)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "16")
        rPr.append(sz)
        run.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        run.append(t)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    # ── title ──
    heading(f"Peer-Review Report", level=0) if False else None
    t = doc.add_heading(r.title or "Review Report", level=0)
    if t.runs:
        t.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    # recommendation banner
    rec_color = _REC_COLOR.get(r.recommendation, (0x33, 0x41, 0x55))
    p = doc.add_paragraph()
    run = p.add_run(f"Recommendation: {r.recommendation}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(*rec_color)
    run2 = p.add_run(f"      Overall score: {r.overall_score:.1f} / 10")
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    if r.verdict:
        para(r.verdict, italic=True, color=(0x47, 0x55, 0x69))

    meta = f"Model: {r.model}  •  Depth: {r.depth}"
    if r.venue_id:
        meta += f"  •  Venue: {r.venue_id}"
    meta += f"  •  Parser: {r.parse_engine}"
    para(meta, size=8.5, color=(0x94, 0xA3, 0xB8))
    doc.add_paragraph()

    if r.headline_summary:
        heading("Summary", level=1)
        doc.add_paragraph(r.headline_summary)

    # scorecard table
    heading("Scorecard", level=1)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Light Grid Accent 1"
    h = tbl.rows[0].cells
    h[0].text, h[1].text = "Dimension", "Score / 10"
    for c in h:
        for rr in c.paragraphs[0].runs:
            rr.bold = True
    for d in r.dimensions:
        row = tbl.add_row().cells
        row[0].text = d.label
        row[1].text = f"{d.score:.1f}"
    counts = r.counts_by_severity()
    para(
        f"Findings: {counts['CRITICAL']} critical · {counts['MAJOR']} major · "
        f"{counts['MINOR']} minor · {counts['PRAISE']} strengths",
        size=9.5, color=(0x47, 0x55, 0x69),
    )

    if r.strengths:
        heading("Strengths", level=1)
        for s in r.strengths:
            doc.add_paragraph(s, style="List Bullet")
    if r.weaknesses:
        heading("Key Weaknesses", level=1)
        for w in r.weaknesses:
            doc.add_paragraph(w, style="List Bullet")

    if r.action_items:
        heading("Action Plan (prioritized)", level=1)
        for it in r.action_items:
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(f"[{it.priority}] ")
            run.bold = True
            col = _SEV_COLOR.get(it.priority, ((0x33, 0x41, 0x55), ""))[0]
            run.font.color.rgb = RGBColor(*col)
            p.add_run(it.text)

    # venue fit
    if r.venue_fit:
        v = r.venue_fit
        heading(f"Venue Fit — {v.venue_name}", level=1)
        para(f"Fit score: {v.fit_score}/100", bold=True)
        if v.scope_rationale:
            doc.add_paragraph(v.scope_rationale)
        if v.rejection_risks:
            para("Top rejection risks:", bold=True)
            for risk in v.rejection_risks:
                doc.add_paragraph(risk, style="List Bullet")
        if v.checklist:
            lbl = "Submission checklist" + (" (model-generated)" if v.checklist_generated else "")
            para(lbl, bold=True)
            for c in v.checklist:
                p = doc.add_paragraph(style="List Bullet")
                mark = p.add_run(f"[{_STATUS_MARK.get(c.status, c.status)}] ")
                mark.bold = True
                p.add_run(c.text)
                if c.note:
                    para(c.note, italic=True, size=9, color=(0x64, 0x74, 0x8B), indent=0.5)

    # citations
    if r.citation_check and r.citation_check.enabled:
        cc = r.citation_check
        heading("Citation Verification", level=1)
        para(f"Checked {cc.n_refs} references · {cc.n_verified} verified · "
             f"{cc.n_suspicious} need attention", bold=True)
        if cc.note:
            para(cc.note, italic=True, color=(0x64, 0x74, 0x8B))
        sus = [c for c in cc.items if c.status in ("NOT_FOUND", "MISMATCH")]
        if sus:
            para("References needing attention:", bold=True)
            for c in sus:
                p = doc.add_paragraph(style="List Bullet")
                mk = p.add_run(f"[{c.status}] ")
                mk.bold = True
                mk.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
                p.add_run(c.title or c.raw[:140])
                if c.note:
                    para(c.note, italic=True, size=9, color=(0x64, 0x74, 0x8B), indent=0.5)

    # detailed findings
    heading("Detailed Findings", level=1)
    for d in r.dimensions:
        heading(f"{d.label} — {d.score:.1f}/10", level=2)
        if d.error:
            para(f"Could not evaluate: {d.error}", italic=True, color=(0xB9, 0x1C, 0x1C))
            continue
        if d.summary:
            doc.add_paragraph(d.summary)
        for f in d.findings:
            color, _ = _SEV_COLOR.get(f.severity, ((0x33, 0x41, 0x55), ""))
            p = doc.add_paragraph()
            run = p.add_run(f"[{f.severity}] {f.title}")
            run.bold = True
            run.font.color.rgb = RGBColor(*color)
            if f.evidence:
                shaded(f'"{f.evidence[:500]}"')
            if f.issue:
                para(f.issue, size=10, indent=0.0)
            if f.recommendation:
                p2 = doc.add_paragraph()
                rb = p2.add_run("Fix: ")
                rb.bold = True
                p2.add_run(f.recommendation)

    if r.warnings:
        heading("Notes", level=1)
        for w in r.warnings:
            doc.add_paragraph(w, style="List Bullet")

    # footer line
    doc.add_paragraph()
    p = doc.add_paragraph()
    lead = p.add_run(f"{GENERATED_BY}  {SEP}  GitHub: ")
    lead.italic = True
    lead.font.size = Pt(8)
    lead.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    add_hyperlink(p, PAIPER_REPO_LABEL, PAIPER_REPO_URL)
    tail = p.add_run(f"  {SEP}  {usage_suffix(r)}")
    tail.italic = True
    tail.font.size = Pt(8)
    tail.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
